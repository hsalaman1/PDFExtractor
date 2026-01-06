"""Word document (.docx) extraction module."""

import hashlib
from datetime import datetime
from pathlib import Path
from typing import Optional
import tempfile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def compute_file_hash(file_path: Path) -> str:
    """Compute SHA256 hash of a file."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return f"sha256:{sha256_hash.hexdigest()}"


def compute_data_hash(data: bytes) -> str:
    """Compute SHA256 hash of data."""
    return f"sha256:{hashlib.sha256(data).hexdigest()}"


def extract_metadata(doc: Document, filename: str) -> dict:
    """Extract metadata from Word document."""
    core_props = doc.core_properties

    return {
        "title": core_props.title or Path(filename).stem,
        "author": core_props.author or None,
        "subject": core_props.subject or None,
        "keywords": core_props.keywords.split(",") if core_props.keywords else [],
        "category": core_props.category or None,
        "comments": core_props.comments or None,
        "created": core_props.created.isoformat() if core_props.created else None,
        "modified": core_props.modified.isoformat() if core_props.modified else None,
        "last_modified_by": core_props.last_modified_by or None,
    }


def extract_toc_from_headings(doc: Document) -> list:
    """Extract table of contents from heading styles."""
    toc = []
    paragraph_num = 0

    for para in doc.paragraphs:
        paragraph_num += 1
        style_name = para.style.name if para.style else ""

        # Check for heading styles
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
                if para.text.strip():
                    toc.append({
                        "level": level,
                        "title": para.text.strip(),
                        "paragraph": paragraph_num
                    })
            except ValueError:
                pass

    return toc


def extract_text_by_section(doc: Document, toc: list) -> list:
    """Extract text organized by sections based on headings."""
    if not toc:
        # No headings - return all content as one section
        full_text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
        return [{
            "section": 1,
            "title": "Content",
            "text": full_text
        }]

    sections = []
    paragraphs = list(doc.paragraphs)
    toc_paragraphs = [entry["paragraph"] for entry in toc]

    for i, entry in enumerate(toc):
        start_idx = entry["paragraph"] - 1
        if i + 1 < len(toc):
            end_idx = toc[i + 1]["paragraph"] - 1
        else:
            end_idx = len(paragraphs)

        section_paras = paragraphs[start_idx:end_idx]
        section_text = "\n\n".join(
            para.text for para in section_paras if para.text.strip()
        )

        sections.append({
            "section": i + 1,
            "title": entry["title"],
            "level": entry["level"],
            "text": section_text
        })

    return sections


def estimate_pages(doc: Document) -> int:
    """Estimate page count based on content."""
    total_chars = sum(len(para.text) for para in doc.paragraphs)
    chars_per_page = 3000
    return max(1, total_chars // chars_per_page + 1)


def extract_docx(
    docx_path: str | Path,
    output_dir: Optional[str | Path] = None,
    output_format: str = "txt",
) -> dict:
    """
    Extract content from a Word document.

    Args:
        docx_path: Path to the Word document
        output_dir: Directory to save extracted content
        output_format: Output format - 'txt', 'json', or 'both'

    Returns:
        Dictionary with extraction results
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"Word document not found: {docx_path}")

    if output_dir is None:
        output_dir = docx_path.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    base_name = docx_path.stem
    file_hash = compute_file_hash(docx_path)

    try:
        doc = Document(docx_path)
    except PackageNotFoundError:
        raise ValueError(f"Invalid or corrupted Word document: {docx_path}")

    # Extract content
    metadata = extract_metadata(doc, docx_path.name)
    toc = extract_toc_from_headings(doc)
    sections = extract_text_by_section(doc, toc)

    # Calculate stats
    total_text = " ".join(s["text"] for s in sections)
    metadata["page_count"] = estimate_pages(doc)
    metadata["word_count"] = len(total_text.split())
    metadata["char_count"] = len(total_text)
    metadata["paragraph_count"] = len(doc.paragraphs)

    result = {
        "source_file": str(docx_path.name),
        "source_path": str(docx_path.absolute()),
        "file_type": "docx",
        "extraction_date": datetime.now().isoformat(),
        "file_hash": file_hash,
        "metadata": metadata,
        "toc": toc,
        "sections": sections,
        "output_files": []
    }

    # Generate text output
    if output_format in ("txt", "both"):
        txt_path = output_dir / f"{base_name}.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"# {metadata.get('title', base_name)}\n")
            if metadata.get("author"):
                f.write(f"# Author: {metadata['author']}\n")
            f.write(f"# Source: {docx_path.name}\n")
            f.write(f"# Words: {metadata['word_count']}\n")
            f.write("=" * 60 + "\n\n")

            for section in sections:
                f.write(f"\n--- Section {section['section']}: {section['title']} ---\n\n")
                f.write(section["text"])
                f.write("\n")

        result["output_files"].append(str(txt_path))
        result["text_output"] = open(txt_path, "r", encoding="utf-8").read()

    # Save JSON output
    if output_format in ("json", "both"):
        import json
        json_path = output_dir / f"{base_name}.json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
        result["output_files"].append(str(json_path))

    return result


def extract_docx_from_data(
    data: bytes,
    filename: str,
    output_format: str = "txt"
) -> dict:
    """Extract content from Word document data (for API use)."""
    file_hash = compute_data_hash(data)

    # Write to temp file for python-docx
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)

        metadata = extract_metadata(doc, filename)
        toc = extract_toc_from_headings(doc)
        sections = extract_text_by_section(doc, toc)

        # Calculate stats
        total_text = " ".join(s["text"] for s in sections)
        metadata["page_count"] = estimate_pages(doc)
        metadata["word_count"] = len(total_text.split())
        metadata["char_count"] = len(total_text)
        metadata["paragraph_count"] = len(doc.paragraphs)

    finally:
        Path(tmp_path).unlink(missing_ok=True)

    result = {
        "source_file": filename,
        "file_type": "docx",
        "extraction_date": datetime.now().isoformat(),
        "file_hash": file_hash,
        "metadata": metadata,
        "toc": toc,
        "sections": sections,
    }

    # Generate text output
    if output_format in ("txt", "both"):
        lines = [
            f"# {metadata.get('title', filename)}",
        ]
        if metadata.get("author"):
            lines.append(f"# Author: {metadata['author']}")
        lines.append(f"# Source: {filename}")
        lines.append(f"# Words: {metadata['word_count']}")
        lines.append("=" * 60)
        lines.append("")

        for section in sections:
            lines.append(f"\n--- Section {section['section']}: {section['title']} ---\n")
            lines.append(section["text"])

        result["text_output"] = "\n".join(lines)

    return result
