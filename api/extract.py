"""Vercel Serverless Function for document extraction.

Supports: PDF, Markdown (.md), Word (.docx)
"""

import base64
import hashlib
import json
import re
import tempfile
from datetime import datetime
from http.server import BaseHTTPRequestHandler
from pathlib import Path

import fitz  # PyMuPDF
import yaml
from docx import Document


def compute_file_hash(data: bytes) -> str:
    """Compute SHA256 hash of file data."""
    return f"sha256:{hashlib.sha256(data).hexdigest()}"


def get_file_type(filename: str) -> str:
    """Determine file type from extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    elif ext in (".md", ".markdown"):
        return "markdown"
    elif ext == ".docx":
        return "docx"
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ============================================================================
# PDF Extraction
# ============================================================================

def extract_pdf_metadata(doc: fitz.Document, filename: str) -> dict:
    """Extract metadata from PDF document."""
    meta = doc.metadata or {}
    keywords = []
    if meta.get("keywords"):
        keywords = [k.strip() for k in meta["keywords"].split(",") if k.strip()]

    return {
        "title": meta.get("title") or Path(filename).stem,
        "author": meta.get("author") or None,
        "subject": meta.get("subject") or None,
        "keywords": keywords,
        "creator": meta.get("creator") or None,
        "creation_date": meta.get("creationDate") or None,
        "page_count": doc.page_count,
    }


def extract_pdf_toc(doc: fitz.Document) -> list:
    """Extract table of contents from PDF."""
    toc = doc.get_toc()
    return [
        {"level": level, "title": title, "page": page}
        for level, title, page in toc
    ]


def extract_pdf_pages(doc: fitz.Document) -> list:
    """Extract text from each page."""
    pages = []
    for page_num in range(doc.page_count):
        page = doc[page_num]
        text = page.get_text("text")
        pages.append({
            "page": page_num + 1,
            "text": text.strip()
        })
    return pages


def extract_pdf(data: bytes, filename: str, output_format: str = "txt") -> dict:
    """Extract content from PDF data."""
    file_hash = compute_file_hash(data)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    try:
        doc = fitz.open(tmp_path)
        metadata = extract_pdf_metadata(doc, filename)
        toc = extract_pdf_toc(doc)
        pages = extract_pdf_pages(doc)
        doc.close()
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    result = {
        "source_file": filename,
        "file_type": "pdf",
        "extraction_date": datetime.now().isoformat(),
        "file_hash": file_hash,
        "metadata": metadata,
        "toc": toc,
        "pages": pages,
    }

    if output_format in ("txt", "both"):
        lines = [f"# {metadata.get('title', filename)}"]
        if metadata.get("author"):
            lines.append(f"# Author: {metadata['author']}")
        lines.append(f"# Source: {filename}")
        lines.append(f"# Pages: {metadata['page_count']}")
        lines.append("=" * 60)
        lines.append("")

        for page_data in pages:
            lines.append(f"\n--- Page {page_data['page']} ---\n")
            lines.append(page_data["text"])

        result["text_output"] = "\n".join(lines)

    return result


# ============================================================================
# Markdown Extraction
# ============================================================================

def extract_markdown_frontmatter(content: str) -> tuple[dict, str]:
    """Extract YAML frontmatter from markdown content."""
    frontmatter = {}
    body = content

    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            try:
                frontmatter = yaml.safe_load(parts[1]) or {}
                body = parts[2].strip()
            except yaml.YAMLError:
                pass

    return frontmatter, body


def extract_markdown_headings(content: str) -> list:
    """Extract headings from markdown to build TOC."""
    toc = []
    lines = content.split("\n")

    for i, line in enumerate(lines):
        match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
        if match:
            level = len(match.group(1))
            title = match.group(2).strip()
            title = re.sub(r"\s*#+\s*$", "", title)
            toc.append({"level": level, "title": title, "line": i + 1})

    return toc


def extract_markdown(data: bytes, filename: str, output_format: str = "txt") -> dict:
    """Extract content from Markdown data."""
    content = data.decode("utf-8")
    file_hash = compute_file_hash(data)

    frontmatter, body = extract_markdown_frontmatter(content)
    toc = extract_markdown_headings(body)

    metadata = {
        "title": frontmatter.get("title") or Path(filename).stem,
        "author": frontmatter.get("author") or frontmatter.get("authors"),
        "subject": frontmatter.get("subject") or frontmatter.get("description"),
        "keywords": frontmatter.get("keywords") or frontmatter.get("tags") or [],
        "date": frontmatter.get("date"),
        "word_count": len(body.split()),
        "char_count": len(body),
    }

    # Split into sections
    sections = []
    if not toc:
        sections = [{"section": 1, "title": "Content", "text": body.strip()}]
    else:
        lines = body.split("\n")
        for i, entry in enumerate(toc):
            start_line = entry["line"] - 1
            end_line = toc[i + 1]["line"] - 1 if i + 1 < len(toc) else len(lines)
            section_text = "\n".join(lines[start_line:end_line]).strip()
            sections.append({
                "section": i + 1,
                "title": entry["title"],
                "level": entry["level"],
                "text": section_text
            })

    result = {
        "source_file": filename,
        "file_type": "markdown",
        "extraction_date": datetime.now().isoformat(),
        "file_hash": file_hash,
        "metadata": metadata,
        "frontmatter": frontmatter,
        "toc": toc,
        "sections": sections,
    }

    if output_format in ("txt", "both"):
        lines = [f"# {metadata.get('title', filename)}"]
        if metadata.get("author"):
            author = metadata["author"]
            if isinstance(author, list):
                author = ", ".join(author)
            lines.append(f"# Author: {author}")
        lines.append(f"# Source: {filename}")
        lines.append(f"# Words: {metadata['word_count']}")
        lines.append("=" * 60)
        lines.append("")

        for section in sections:
            lines.append(f"\n--- Section {section['section']}: {section['title']} ---\n")
            lines.append(section["text"])

        result["text_output"] = "\n".join(lines)

    return result


# ============================================================================
# Word Document Extraction
# ============================================================================

def extract_docx_metadata(doc: Document, filename: str) -> dict:
    """Extract metadata from Word document."""
    core_props = doc.core_properties

    return {
        "title": core_props.title or Path(filename).stem,
        "author": core_props.author or None,
        "subject": core_props.subject or None,
        "keywords": core_props.keywords.split(",") if core_props.keywords else [],
        "category": core_props.category or None,
        "created": core_props.created.isoformat() if core_props.created else None,
        "modified": core_props.modified.isoformat() if core_props.modified else None,
    }


def extract_docx_toc(doc: Document) -> list:
    """Extract table of contents from heading styles."""
    toc = []
    paragraph_num = 0

    for para in doc.paragraphs:
        paragraph_num += 1
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
                if para.text.strip():
                    toc.append({
                        "level": level,
                        "title": para.text.strip(),
                        "paragraph": paragraph_num
                    })
            except ValueError:
                pass

    return toc


def extract_docx(data: bytes, filename: str, output_format: str = "txt") -> dict:
    """Extract content from Word document data."""
    file_hash = compute_file_hash(data)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        metadata = extract_docx_metadata(doc, filename)
        toc = extract_docx_toc(doc)

        # Extract sections
        sections = []
        paragraphs = list(doc.paragraphs)

        if not toc:
            full_text = "\n\n".join(para.text for para in paragraphs if para.text.strip())
            sections = [{"section": 1, "title": "Content", "text": full_text}]
        else:
            for i, entry in enumerate(toc):
                start_idx = entry["paragraph"] - 1
                end_idx = toc[i + 1]["paragraph"] - 1 if i + 1 < len(toc) else len(paragraphs)
                section_paras = paragraphs[start_idx:end_idx]
                section_text = "\n\n".join(para.text for para in section_paras if para.text.strip())
                sections.append({
                    "section": i + 1,
                    "title": entry["title"],
                    "level": entry["level"],
                    "text": section_text
                })

        # Calculate stats
        total_text = " ".join(s["text"] for s in sections)
        metadata["word_count"] = len(total_text.split())
        metadata["char_count"] = len(total_text)
        metadata["paragraph_count"] = len(paragraphs)

    finally:
        Path(tmp_path).unlink(missing_ok=True)

    result = {
        "source_file": filename,
        "file_type": "docx",
        "extraction_date": datetime.now().isoformat(),
        "file_hash": file_hash,
        "metadata": metadata,
        "toc": toc,
        "sections": sections,
    }

    if output_format in ("txt", "both"):
        lines = [f"# {metadata.get('title', filename)}"]
        if metadata.get("author"):
            lines.append(f"# Author: {metadata['author']}")
        lines.append(f"# Source: {filename}")
        lines.append(f"# Words: {metadata['word_count']}")
        lines.append("=" * 60)
        lines.append("")

        for section in sections:
            lines.append(f"\n--- Section {section['section']}: {section['title']} ---\n")
            lines.append(section["text"])

        result["text_output"] = "\n".join(lines)

    return result


# ============================================================================
# HTTP Handler
# ============================================================================

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers.get("Content-Length", 0))
            content_type = self.headers.get("Content-Type", "")

            if content_length == 0:
                self.send_error(400, "No data received")
                return

            body = self.rfile.read(content_length)

            if "application/json" not in content_type:
                self.send_error(400, "Please send JSON with base64 encoded file")
                return

            data = json.loads(body)
            file_data = base64.b64decode(data.get("file", ""))
            filename = data.get("filename", "document.pdf")
            output_format = data.get("format", "txt")

            if not file_data:
                self.send_error(400, "No file data provided")
                return

            # Determine file type and extract
            try:
                file_type = get_file_type(filename)
            except ValueError as e:
                self.send_error(400, str(e))
                return

            if file_type == "pdf":
                result = extract_pdf(file_data, filename, output_format)
            elif file_type == "markdown":
                result = extract_markdown(file_data, filename, output_format)
            elif file_type == "docx":
                result = extract_docx(file_data, filename, output_format)

            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(json.dumps(result).encode())

        except Exception as e:
            self.send_response(500)
            self.send_header("Content-Type", "application/json")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(json.dumps({"error": str(e)}).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()
